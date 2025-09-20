from adapters.bedrock_model_adapter import BedrockModelAdapter
from typing import Dict, Any, Optional, List
import json
from datetime import datetime
import re
import logging

# PDF/Word imports 
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from io import BytesIO

# Database imports
from supabase import Client
import os
from dotenv import load_dotenv

load_dotenv()
logger = logging.getLogger(__name__)

class MedicalReportNode:
    def __init__(self, adapter: BedrockModelAdapter, supabase_client: Optional[Client] = None):
        self.adapter = adapter
        
        # Initialize Supabase client for database operations
        if supabase_client:
            self.supabase = supabase_client
        else:
            # Create client if not provided
            url = os.getenv("SUPABASE_URL")
            key = os.getenv("SUPABASE_API_KEY")
            if url and key:
                from supabase import create_client
                self.supabase = create_client(url, key)
            else:
                self.supabase = None
                logger.warning("Supabase credentials not found - database features disabled")
    
    async def __call__(self, state: dict) -> dict:
        """Generate medical report content and optionally save to database"""
        print("📄 MEDICAL REPORT NODE CALLED!")
        
        # Set stage when node starts
        state["current_workflow_stage"] = "generating_medical_report"
        
        # Generate the actual medical report content
        state = await self.generate_medical_report_content(state)
        
        # Mark as complete
        state["current_workflow_stage"] = "workflow_complete"
        
        return state
    
    async def generate_medical_report_content(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """Generate optimized medical report using templates + minimal LLM generation"""
        
        try:
            print("🔄 Generating template-based medical report...")
            
            # Generate only the dynamic content via LLM
            dynamic_content = await self._generate_followup_guidance(state)
            
            # Combine with template-based sections
            final_report = self._create_template_based_report(state, dynamic_content)
            
            # Store in state
            state["medical_report"] = final_report
            
            print("✅ Template-based medical report generated successfully")
            return state
                
        except Exception as e:
            print(f"❌ Medical report generation failed: {e}")
            # Generate fallback report
            state["medical_report"] = self._generate_fallback_report(state)
            return state

    # ================================
    # DATABASE STORAGE METHODS
    # ================================
    
    async def save_medical_report_to_database(
        self, 
        user_id: str, 
        session_id: str, 
        agent_state: Dict[str, Any],
        report_title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Save a complete medical report to the database"""
        
        if not self.supabase:
            raise Exception("Database not configured - Supabase client not available")
        
        try:
            # Extract data from agent state
            report_data = {
                "user_id": user_id,
                "session_id": session_id,
                "report_title": report_title or f"Medical Report - {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                "patient_symptoms": agent_state.get("userInput_symptoms") or agent_state.get("userInput_skin_symptoms"),
                "textual_analysis": agent_state.get("textual_analysis"),
                "followup_data": {
                    "questions": agent_state.get("followup_questions"),
                    "responses": agent_state.get("followup_response"),
                    "qna_overall": agent_state.get("followup_qna_overall"),
                    "diagnosis": agent_state.get("followup_diagnosis")
                } if agent_state.get("followup_questions") else None,
                "image_analysis": agent_state.get("skin_lesion_analysis"),
                "overall_analysis": agent_state.get("overall_analysis"),
                "healthcare_recommendations": agent_state.get("healthcare_recommendation"),
                "medical_report_content": agent_state.get("medical_report"),
                "workflow_path": agent_state.get("workflow_path"),
                "workflow_stages_completed": agent_state.get("current_workflow_stage"),
                "confidence_scores": {
                    "average_confidence": agent_state.get("average_confidence"),
                    "final_confidence": agent_state.get("overall_analysis", {}).get("final_confidence") if agent_state.get("overall_analysis") else None
                }
            }
            
            # Remove None values
            report_data = {k: v for k, v in report_data.items() if v is not None}
            
            # Insert into database
            result = self.supabase.table("medical_reports").insert(report_data).execute()
            
            if result.data:
                logger.info(f"Medical report saved successfully for user {user_id}, session {session_id}")
                return result.data[0]
            else:
                raise Exception("Failed to save medical report")
                
        except Exception as e:
            logger.error(f"Error saving medical report: {e}")
            raise e
    
    async def get_user_medical_reports(
        self, 
        user_id: str, 
        limit: int = 10, 
        offset: int = 0
    ) -> List[Dict[str, Any]]:
        """Get all medical reports for a user"""
        
        if not self.supabase:
            raise Exception("Database not configured")
        
        try:
            result = self.supabase.table("medical_reports")\
                .select("*")\
                .eq("user_id", user_id)\
                .order("created_at", desc=True)\
                .range(offset, offset + limit - 1)\
                .execute()
            
            return result.data or []
            
        except Exception as e:
            logger.error(f"Error fetching medical reports: {e}")
            raise e
    
    async def get_medical_report_by_id(
        self, 
        report_id: str, 
        user_id: str
    ) -> Optional[Dict[str, Any]]:
        """Get a specific medical report by ID"""
        
        if not self.supabase:
            raise Exception("Database not configured")
        
        try:
            result = self.supabase.table("medical_reports")\
                .select("*")\
                .eq("id", report_id)\
                .eq("user_id", user_id)\
                .execute()
            
            return result.data[0] if result.data else None
            
        except Exception as e:
            logger.error(f"Error fetching medical report: {e}")
            raise e
    
    async def delete_medical_report(
        self, 
        report_id: str, 
        user_id: str
    ) -> bool:
        """Delete a medical report"""
        
        if not self.supabase:
            raise Exception("Database not configured")
        
        try:
            result = self.supabase.table("medical_reports")\
                .delete()\
                .eq("id", report_id)\
                .eq("user_id", user_id)\
                .execute()
            
            return bool(result.data)
            
        except Exception as e:
            logger.error(f"Error deleting medical report: {e}")
            raise e

    async def update_report_title(
        self, 
        report_id: str, 
        user_id: str, 
        new_title: str
    ) -> Dict[str, Any]:
        """Update medical report title"""
        
        if not self.supabase:
            raise Exception("Database not configured")
        
        try:
            result = self.supabase.table("medical_reports")\
                .update({"report_title": new_title})\
                .eq("id", report_id)\
                .eq("user_id", user_id)\
                .execute()
            
            return result.data[0] if result.data else None
            
        except Exception as e:
            logger.error(f"Error updating medical report title: {e}")
            raise e

    # ================================
    # EXISTING METHODS (unchanged)
    # ================================
    
    async def _generate_followup_guidance(self, state: Dict[str, Any]) -> Dict[str, str]:
        """Generate only the content that needs LLM creativity"""
        
        overall_analysis = state.get("overall_analysis", {})
        diagnosis = overall_analysis.get('final_diagnosis', 'Unknown')
        severity = overall_analysis.get('final_severity', 'moderate')
        confidence = overall_analysis.get('final_confidence', 0.5)
        specialist = overall_analysis.get('specialist_recommendation', 'general_practitioner')
        
        # Short, focused prompts for specific content
        followup_prompt = f"""Generate follow-up guidance for:
    Diagnosis: {diagnosis}
    Severity: {severity}
    Confidence: {confidence:.2f}
    Specialist: {specialist}

    The patient has NOT yet seen the specialist. provide in this EXACT format:
    IMMEDIATE (24-48h): [scheduling specialist appointment and urgent self-care]
    SHORT-TERM (1-2 weeks): [monitoring/appointments]
    WATCH FOR: [warning signs]
    LIFESTYLE: [safe modifications until specialist consultation]

    Keep each section under 30 words."""
    
        # Generate both sections concurrently for speed
        followup_guidance = await self.adapter.generate_text_guidance(followup_prompt, 200, 0.2)
        
        return {
            "followup_guidance": followup_guidance,
        }
    
    def _create_template_based_report(self, state: Dict[str, Any], dynamic_content: Dict[str, str]) -> str:
        """Create report using templates filled with state data + dynamic content"""
        
        overall_analysis = state.get("overall_analysis", {})
        session_id = state.get("session_id", "Unknown")
        workflow_path = state.get("workflow_path", [])
        
        # Extract data for templates
        diagnosis = overall_analysis.get('final_diagnosis', 'Not determined')
        confidence = overall_analysis.get('final_confidence', 0.0)
        severity = overall_analysis.get('final_severity', 'moderate')
        specialist = overall_analysis.get('specialist_recommendation', 'general_practitioner')
        user_explanation = overall_analysis.get('user_explanation', 'Analysis completed')
        clinical_reasoning = overall_analysis.get('clinical_reasoning', 'Systematic analysis performed')
        
        # Analysis type from workflow
        analysis_type = self._get_analysis_type_display(workflow_path, state)
        
        # Get alternative diagnoses
        alternative_diagnoses = self._get_alternative_diagnoses(state)
        
        # Urgency level based on severity
        urgency_info = self._get_urgency_template(severity)
        
        # Build complete report
        report = f"""
    MEDICAL ANALYSIS REPORT
    ═══════════════════════════════════════════════════════════════════
    Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}
    Session ID: {session_id}
    Analysis Method: {analysis_type}

═══════════════════════════════════════════════════════════════════
        EXECUTIVE SUMMARY
═══════════════════════════════════════════════════════════════════
    Primary Diagnosis: {diagnosis}
    Diagnostic Confidence: {(confidence * 100):.1f}%
    Severity Level: {severity.title()}
    Recommended Specialist: {specialist.replace('_', ' ').title()}

    {urgency_info['summary']}

═══════════════════════════════════════════════════════════════════
        AI CLINICAL ASSESSMENT
═══════════════════════════════════════════════════════════════════
    CONDITION OVERVIEW:
    {user_explanation}

    CLINICAL REASONING:
    {clinical_reasoning}

    DIAGNOSTIC CONFIDENCE:
    Based on the available evidence, this diagnosis has a confidence level of {(confidence * 100):.1f}%. 
    {self._get_confidence_interpretation(confidence)}

    ALTERNATIVE DIAGNOSES:
    {alternative_diagnoses}

═══════════════════════════════════════════════════════════════════
        FOLLOW-UP GUIDANCE
═══════════════════════════════════════════════════════════════════
    {dynamic_content.get('followup_guidance', 'Follow standard care protocols for this condition.')}

═══════════════════════════════════════════════════════════════════
        SPECIALIST REFERRAL
═══════════════════════════════════════════════════════════════════
    RECOMMENDED SPECIALIST: {specialist.replace('_', ' ').title()}

    REFERRAL TIMING: {self._get_referral_timing(severity)}

    PREPARATION FOR APPOINTMENT:
    • Bring this report and any previous medical records
    • List all current medications and supplements
    • Prepare a detailed symptom timeline
    • Note any family history of similar conditions
    • Bring insurance information and identification

═══════════════════════════════════════════════════════════════════
        SAFETY WARNINGS
═══════════════════════════════════════════════════════════════════
    {urgency_info['warnings']}

    GENERAL EMERGENCY SIGNS - Seek immediate medical attention for:
    • Severe difficulty breathing or shortness of breath
    • Chest pain or pressure lasting more than a few minutes
    • Severe bleeding that won't stop
    • Loss of consciousness or severe confusion
    • Signs of severe allergic reaction (swelling, difficulty swallowing)
    • Sudden severe headache with vision changes
    • High fever (over 103°F/39.4°C) with severe symptoms

═══════════════════════════════════════════════════════════════════
        MEDICAL DISCLAIMERS
═══════════════════════════════════════════════════════════════════
    AI ANALYSIS LIMITATIONS:
    This report is generated by an AI medical analysis system and is intended for 
    informational and educational purposes only. It does not constitute professional 
    medical advice, diagnosis, or treatment recommendations.

    PROFESSIONAL CONSULTATION REQUIRED:
    • This analysis cannot replace professional medical examination
    • Physical examination, laboratory tests, and imaging may be necessary
    • A qualified healthcare provider should review all symptoms
    • Treatment decisions should only be made by licensed medical professionals

    ACCURACY CONSIDERATIONS:
    • Analysis accuracy depends on completeness of provided information
    • Some conditions require specialized testing for proper diagnosis
    • AI systems may not detect all possible conditions or complications
    • Second medical opinions are recommended for complex cases

    EMERGENCY DISCLAIMER:
    If you are experiencing a medical emergency, do not rely on this report. 
    Call emergency services (911) or go to the nearest emergency room immediately.

    DATA PRIVACY:
    This analysis is confidential and should be shared only with your healthcare 
    providers. Maintain privacy of medical information according to applicable laws.

    ═══════════════════════════════════════════════════════════════════
    Report ID: {session_id} | AI Medical Analysis System v2.0
    Total Analysis Time: {self._get_analysis_duration(state)}
    Generated by Llama 3.1 8B UltraMedical (8-bit Quantization GGUF)
    ═══════════════════════════════════════════════════════════════════
    """
        
        return report
    
    def _generate_fallback_report(self, state: Dict[str, Any]) -> str:
        """Generate a fallback report if LLM generation fails"""
        
        overall_analysis = state.get("overall_analysis", {})
        session_id = state.get("session_id", "Unknown")
        
        return f"""
MEDICAL ANALYSIS REPORT
Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}
Session ID: {session_id}

═══════════════════════════════════════════════════════════════════
            EXECUTIVE SUMMARY
═══════════════════════════════════════════════════════════════════

Primary Diagnosis: {overall_analysis.get('final_diagnosis', 'Analysis incomplete')}

Based on the comprehensive analysis of provided symptoms, our AI medical assistant has completed the diagnostic assessment with the findings detailed below.

═══════════════════════════════════════════════════════════════════
            CLINICAL FINDINGS
═══════════════════════════════════════════════════════════════════

DIAGNOSTIC ASSESSMENT:
• Primary Diagnosis: {overall_analysis.get('final_diagnosis', 'Not determined')}
• Confidence Level: {(overall_analysis.get('final_confidence', 0) * 100):.1f}%
• Severity Assessment: {overall_analysis.get('final_severity', 'unknown').title()}

CLINICAL EXPLANATION:
{overall_analysis.get('user_explanation', 'Comprehensive analysis was performed based on the provided symptom information.')}

CLINICAL REASONING:
{overall_analysis.get('clinical_reasoning', 'Diagnosis determined through systematic analysis of symptoms and clinical indicators.')}

═══════════════════════════════════════════════════════════════════
            RECOMMENDATIONS
═══════════════════════════════════════════════════════════════════

SPECIALIST CONSULTATION:
Recommended Specialist: {overall_analysis.get('specialist_recommendation', 'general_practitioner').replace('_', ' ').title()}

NEXT STEPS:
1. Schedule consultation with recommended specialist
2. Continue monitoring symptoms as described
3. Seek immediate medical attention if symptoms worsen
4. Follow up as directed by healthcare provider

═══════════════════════════════════════════════════════════════════
            IMPORTANT DISCLAIMERS
═══════════════════════════════════════════════════════════════════

MEDICAL DISCLAIMER:
This AI-generated report is for informational and educational purposes only and should not replace professional medical advice, diagnosis, or treatment. Always consult with qualified healthcare professionals for medical decisions and concerns.

EMERGENCY GUIDANCE:
If you experience severe symptoms, difficulty breathing, chest pain, or other emergency conditions, seek immediate medical attention by calling emergency services.

ACCURACY CONSIDERATIONS:
• Analysis accuracy depends on completeness of provided information
• Some conditions may require physical examination for proper diagnosis
• Follow-up testing may be necessary for definitive diagnosis
• Second opinions are recommended for complex medical cases

═══════════════════════════════════════════════════════════════════
Report generated by Llama 3.1 8B UltraMedical (8-bit Quantization GGUF)
Session: {session_id} | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
═══════════════════════════════════════════════════════════════════
"""

    def _get_alternative_diagnoses(self, state: Dict[str, Any]) -> str:
        """Extract and format alternative diagnoses from state data"""
        
        alternative_diagnoses = []
        
        # Get diagnoses based on workflow path
        textual_analysis = state.get("textual_analysis", [])
        followup_diagnosis = state.get("followup_diagnosis", [])
        skin_lesion_analysis = state.get("skin_lesion_analysis", {})
        
        # Determine which diagnosis list to use
        if followup_diagnosis and len(followup_diagnosis) > 1:
            # Use enhanced follow-up diagnoses
            primary_diagnosis = followup_diagnosis[0].get("text_diagnosis", "")
            for i, diag in enumerate(followup_diagnosis[1:], 1):
                if i <= 3:  # Show top 3 alternatives
                    diagnosis_text = diag.get("text_diagnosis", "Unknown")
                    confidence = diag.get("diagnosis_confidence", 0.0)
                    alternative_diagnoses.append(f"{i}. {diagnosis_text} ({confidence:.1%} confidence)")
        
        elif textual_analysis and len(textual_analysis) > 1:
            # Use initial textual analysis alternatives
            primary_diagnosis = textual_analysis[0].get("text_diagnosis", "")
            for i, diag in enumerate(textual_analysis[1:], 1):
                if i <= 3:  # Show top 3 alternatives
                    diagnosis_text = diag.get("text_diagnosis", "Unknown")
                    confidence = diag.get("diagnosis_confidence", 0.0)
                    alternative_diagnoses.append(f"{i}. {diagnosis_text} ({confidence:.1%} confidence)")
        
        elif skin_lesion_analysis.get("confidence_score"):
            # Use image analysis alternatives for skin conditions
            confidence_scores = skin_lesion_analysis.get("confidence_score", {})
            sorted_conditions = sorted(confidence_scores.items(), key=lambda x: x[1], reverse=True)
            
            # Skip the primary diagnosis (highest confidence) and show alternatives
            for i, (condition, confidence) in enumerate(sorted_conditions[1:4], 1):
                alternative_diagnoses.append(f"{i}. {condition} ({confidence:.1f}% confidence)")
        
        if alternative_diagnoses:
            alternatives_text = "\n".join([f"• {alt}" for alt in alternative_diagnoses])
            return f"The following alternative diagnoses were also considered:\n{alternatives_text}\n\nThese alternatives may warrant further evaluation if primary diagnosis is ruled out."
        else:
            return "No significant alternative diagnoses identified based on current analysis. Additional testing may reveal other possibilities."

    def _get_analysis_type_display(self, workflow_path: list, state: Dict[str, Any]) -> str:
        """Get user-friendly analysis type description"""
        has_followup = bool(state.get("followup_response"))
        has_image = bool(state.get("skin_lesion_analysis", {}).get("image_diagnosis"))
        
        if has_followup and has_image:
            return "Comprehensive Multi-Modal Analysis (Symptoms + Follow-up + Image)"
        elif has_image:
            return "Visual + Symptom Analysis (Dermatological Screening)"
        elif has_followup:
            return "Enhanced Symptom Analysis (with Follow-up Questions)"
        else:
            return "Standard Symptom Analysis"

    def _get_urgency_template(self, severity: str) -> Dict[str, str]:
        """Get urgency information based on severity"""
        urgency_templates = {
            "critical": {
                "summary": "⚠️ URGENT: This condition requires immediate medical attention. Seek emergency care or contact your healthcare provider immediately.",
                "warnings": "IMMEDIATE ACTION REQUIRED:\n• Contact emergency services if experiencing severe symptoms\n• Do not delay seeking professional medical care\n• Monitor symptoms closely and seek help if they worsen"
            },
            "severe": {
                "summary": "⚠️ HIGH PRIORITY: This condition requires prompt medical evaluation within 24-48 hours.",
                "warnings": "PROMPT CARE NEEDED:\n• Schedule urgent appointment with healthcare provider\n• Monitor symptoms closely for any worsening\n• Seek emergency care if symptoms become severe"
            },
            "moderate": {
                "summary": "📋 MODERATE PRIORITY: Schedule medical consultation within 1-2 weeks for proper evaluation.",
                "warnings": "MONITORING REQUIRED:\n• Schedule appointment with healthcare provider\n• Watch for symptom progression or new symptoms\n• Seek prompt care if condition worsens"
            },
            "mild": {
                "summary": "📝 ROUTINE FOLLOW-UP: Consider scheduling regular check-up with healthcare provider.",
                "warnings": "GENERAL MONITORING:\n• Continue observing symptoms\n• Schedule routine follow-up as appropriate\n• Contact healthcare provider if symptoms persist or worsen"
            }
        }
        
        return urgency_templates.get(severity.lower(), urgency_templates["moderate"])

    def _get_confidence_interpretation(self, confidence: float) -> str:
        """Get confidence level interpretation"""
        if confidence >= 0.8:
            return "This represents high diagnostic confidence based on clear symptom patterns."
        elif confidence >= 0.6:
            return "This represents moderate diagnostic confidence. Additional evaluation may be helpful."
        elif confidence >= 0.4:
            return "This represents preliminary assessment. Professional evaluation is recommended for confirmation."
        else:
            return "This represents initial screening only. Comprehensive medical evaluation is strongly recommended."

    def _get_referral_timing(self, severity: str) -> str:
        """Get referral timing based on severity"""
        timing_map = {
            "critical": "Immediate - Seek emergency care now",
            "severe": "Urgent - Within 24-48 hours",
            "moderate": "Prompt - Within 1-2 weeks",
            "mild": "Routine - Within 4-6 weeks or as convenient"
        }
        return timing_map.get(severity.lower(), "As recommended by primary care provider")

    def _get_analysis_duration(self, state: Dict[str, Any]) -> str:
        """Get analysis duration if available"""
        # You can track this in your workflow
        return state.get("analysis_duration", "< 2 minutes")

#==================================================
# Medical Report Export Function
#==================================================

    # Export functionality (only called from API endpoint)
    async def generate_export_file(self, state: dict, format: str, include_details: bool = True) -> bytes:
        """Generate PDF or Word export file (separate from main workflow)"""
        
        if format == 'pdf':
            return await self._generate_pdf_export(state, include_details)
        elif format == 'word':
            return await self._generate_word_export(state, include_details)
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    async def _generate_pdf_export(self, state: dict, include_details: bool) -> bytes:
        """Generate PDF using reportlab with template-based content"""
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, 
                            topMargin=72, bottomMargin=18)
            styles = getSampleStyleSheet()
            story = []
            
            # Use the same template-based report generation
            if not state.get("medical_report"):
                # Generate the template-based report if not already generated
                dynamic_content = await self._generate_followup_guidance(state)
                full_report = self._create_template_based_report(state, dynamic_content)
            else:
                full_report = state.get("medical_report", "")
            
            # Extract key information for PDF formatting
            overall_analysis = state.get('overall_analysis', {})
            
            # Title
            story.append(Paragraph("Medical Analysis Report", styles['Title']))
            story.append(Spacer(1, 12))
            
            # Header information
            story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", styles['Normal']))
            story.append(Paragraph(f"Session ID: {state.get('session_id', 'Unknown')}", styles['Normal']))
            story.append(Spacer(1, 12))
            
            if include_details:
                # Use the full template-based report
                # Split the report into sections and format appropriately
                sections = full_report.split('═══════════════════════════════════════════════════════════════════')
                
                for section in sections:
                    if section.strip():
                        # Clean up the section text
                        section_text = section.strip()
                        if section_text:
                            # Check if it's a header (all caps or specific patterns)
                            lines = section_text.split('\n')
                            for line in lines:
                                line = line.strip()
                                if line:
                                    # Detect headers vs content
                                    if line.isupper() or line.startswith('MEDICAL ANALYSIS REPORT') or 'EXECUTIVE SUMMARY' in line:
                                        story.append(Paragraph(line, styles['Heading2']))
                                    elif line.startswith('Generated by') or line.startswith('Report ID'):
                                        story.append(Paragraph(line, styles['Normal']))
                                        story.append(Spacer(1, 6))
                                    else:
                                        # Regular content
                                        if len(line) > 200:  # Long paragraphs
                                            # Split long content into smaller paragraphs
                                            for paragraph in line.split('. '):
                                                if paragraph.strip():
                                                    story.append(Paragraph(paragraph.strip() + '.', styles['Normal']))
                                                    story.append(Spacer(1, 3))
                                        else:
                                            story.append(Paragraph(line, styles['Normal']))
                                            story.append(Spacer(1, 3))
            else:
                # Summary version - extract key information
                story.append(Paragraph("Executive Summary", styles['Heading2']))
                story.append(Paragraph(f"Primary Diagnosis: {overall_analysis.get('final_diagnosis', 'N/A')}", styles['Normal']))
                story.append(Paragraph(f"Confidence Level: {(overall_analysis.get('final_confidence', 0) * 100):.1f}%", styles['Normal']))
                story.append(Paragraph(f"Severity: {overall_analysis.get('final_severity', 'N/A').title()}", styles['Normal']))
                story.append(Paragraph(f"Recommended Specialist: {overall_analysis.get('specialist_recommendation', 'General Practitioner').replace('_', ' ').title()}", styles['Normal']))
                story.append(Spacer(1, 12))
                
                # User explanation
                if overall_analysis.get('user_explanation'):
                    story.append(Paragraph("Condition Overview", styles['Heading2']))
                    story.append(Paragraph(overall_analysis.get('user_explanation'), styles['Normal']))
                    story.append(Spacer(1, 12))
            
            # Always include disclaimer
            story.append(Paragraph("Medical Disclaimer", styles['Heading3']))
            disclaimer_text = ("This AI-generated report is for informational purposes only and should not replace "
                            "professional medical advice, diagnosis, or treatment. Always consult with qualified "
                            "healthcare professionals for medical concerns.")
            story.append(Paragraph(disclaimer_text, styles['Normal']))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"❌ PDF generation failed: {e}")
            # Fallback to text
            return self._generate_text_export(state, include_details).encode('utf-8')

    async def _generate_word_export(self, state: Dict[str, Any], include_details: bool) -> bytes:
        """Generate Word document using template-based content"""
        try:
            doc = Document()
            
            # Use the same template-based report generation
            if not state.get("medical_report"):
                # Generate the template-based report if not already generated
                dynamic_content = await self._generate_followup_guidance(state)
                full_report = self._create_template_based_report(state, dynamic_content)
            else:
                full_report = state.get("medical_report", "")
            
            # Extract key information
            overall_analysis = state.get('overall_analysis', {})
            
            # Title
            title = doc.add_heading('Medical Analysis Report', 0)
            title.alignment = 1  # Center alignment
            
            # Header info
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
            doc.add_paragraph(f"Session ID: {state.get('session_id', 'Unknown')}")
            doc.add_paragraph("")  # Empty line
            
            if include_details:
                # Use the full template-based report
                # Process the full report content
                sections = full_report.split('═══════════════════════════════════════════════════════════════════')
                
                for section in sections:
                    if section.strip():
                        section_text = section.strip()
                        if section_text:
                            lines = section_text.split('\n')
                            current_heading = None
                            current_content = []
                            
                            for line in lines:
                                line = line.strip()
                                if line:
                                    # Detect headers
                                    if (line.isupper() and len(line) < 50) or 'EXECUTIVE SUMMARY' in line or 'CLINICAL ASSESSMENT' in line:
                                        # Add previous content if any
                                        if current_content:
                                            content_text = '\n'.join(current_content)
                                            if content_text.strip():
                                                doc.add_paragraph(content_text)
                                            current_content = []
                                        
                                        # Add new heading
                                        if line != 'MEDICAL ANALYSIS REPORT':
                                            doc.add_heading(line.title(), level=1)
                                        current_heading = line
                                    else:
                                        # Regular content
                                        current_content.append(line)
                            
                            # Add remaining content
                            if current_content:
                                content_text = '\n'.join(current_content)
                                if content_text.strip():
                                    doc.add_paragraph(content_text)
            else:
                # Summary version
                doc.add_heading('Executive Summary', level=1)
                
                # Create a table for key information
                table = doc.add_table(rows=4, cols=2)
                table.style = 'Table Grid'
                
                # Fill table with key information
                cells = table.rows[0].cells
                cells[0].text = 'Primary Diagnosis'
                cells[1].text = overall_analysis.get('final_diagnosis', 'N/A')
                
                cells = table.rows[1].cells
                cells[0].text = 'Confidence Level'
                cells[1].text = f"{(overall_analysis.get('final_confidence', 0) * 100):.1f}%"
                
                cells = table.rows[2].cells
                cells[0].text = 'Severity'
                cells[1].text = overall_analysis.get('final_severity', 'N/A').title()
                
                cells = table.rows[3].cells
                cells[0].text = 'Recommended Specialist'
                cells[1].text = overall_analysis.get('specialist_recommendation', 'General Practitioner').replace('_', ' ').title()
                
                doc.add_paragraph("")  # Empty line
                
                # User explanation
                if overall_analysis.get('user_explanation'):
                    doc.add_heading('Condition Overview', level=1)
                    doc.add_paragraph(overall_analysis.get('user_explanation'))
            
            # Always include disclaimer
            doc.add_heading('Medical Disclaimer', level=2)
            disclaimer_text = ("This AI-generated report is for informational purposes only and should not replace "
                            "professional medical advice, diagnosis, or treatment. Always consult with qualified "
                            "healthcare professionals for medical concerns.")
            doc.add_paragraph(disclaimer_text)
            
            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"❌ Word generation failed: {e}")
            # Fallback to text
            return self._generate_text_export(state, include_details).encode('utf-8')

    def _generate_text_export(self, state: dict, include_details: bool) -> str:
        """Generate plain text export using template-based content"""
        
        # Use the same template-based report generation
        if not state.get("medical_report"):
            # Need to generate the report content
            overall_analysis = state.get("overall_analysis", {})
            
            if include_details:
                # Create a simplified template for text export
                return f"""MEDICAL ANALYSIS REPORT
    Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}
    Session ID: {state.get('session_id', 'Unknown')}

    PRIMARY DIAGNOSIS: {overall_analysis.get('final_diagnosis', 'N/A')}
    CONFIDENCE: {(overall_analysis.get('final_confidence', 0) * 100):.1f}%
    SEVERITY: {overall_analysis.get('final_severity', 'N/A').title()}
    SPECIALIST: {overall_analysis.get('specialist_recommendation', 'General Practitioner').replace('_', ' ').title()}

    CONDITION OVERVIEW:
    {overall_analysis.get('user_explanation', 'Analysis completed based on provided symptoms.')}

    CLINICAL REASONING:
    {overall_analysis.get('clinical_reasoning', 'Diagnosis determined through systematic analysis of symptoms and clinical indicators.')}

    MEDICAL DISCLAIMER:
    This AI-generated report is for informational purposes only and should not replace professional medical advice, diagnosis, or treatment. Always consult with qualified healthcare professionals for medical concerns."""
            else:
                return f"""MEDICAL ANALYSIS REPORT - SUMMARY
    Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}

    PRIMARY DIAGNOSIS: {overall_analysis.get('final_diagnosis', 'N/A')}
    CONFIDENCE: {(overall_analysis.get('final_confidence', 0) * 100):.1f}%
    SEVERITY: {overall_analysis.get('final_severity', 'N/A').title()}
    SPECIALIST: {overall_analysis.get('specialist_recommendation', 'General Practitioner').replace('_', ' ').title()}

    WHAT IS IT?
    {overall_analysis.get('user_explanation', 'Analysis completed based on provided symptoms.')}

    MEDICAL DISCLAIMER:
    This AI-generated report is for informational purposes only and should not replace professional medical advice, diagnosis, or treatment. Always consult with qualified healthcare professionals for medical concerns."""
        else:
            # Use existing template-based report
            full_report = state.get("medical_report", "")
            
            if include_details:
                return full_report
            else:
                # Extract summary from full report
                overall_analysis = state.get("overall_analysis", {})
                return f"""MEDICAL ANALYSIS REPORT - SUMMARY
    Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}

    PRIMARY DIAGNOSIS: {overall_analysis.get('final_diagnosis', 'N/A')}
    CONFIDENCE: {(overall_analysis.get('final_confidence', 0) * 100):.1f}%
    SEVERITY: {overall_analysis.get('final_severity', 'N/A').title()}
    SPECIALIST: {overall_analysis.get('specialist_recommendation', 'General Practitioner').replace('_', ' ').title()}

    WHAT IS IT?
    {overall_analysis.get('user_explanation', 'Analysis completed based on provided symptoms.')}

    MEDICAL DISCLAIMER:
This AI-generated report is for informational purposes only and should not replace professional medical advice, diagnosis, or treatment. Always consult with qualified healthcare professionals for medical concerns."""